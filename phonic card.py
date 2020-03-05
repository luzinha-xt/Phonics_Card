from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
def set_cell_border( _Cell, **kwargs):
    
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = _Cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
"""


five_letter_phonics.docx


"""
# read word list
file_object = open('five_letter_phonics.txt', 'r')
words = file_object.read()
items = list(words)
i = len(items)
print(items, i)
document = Document()

#### set font
style = document.styles['Normal']
font = style.font
font.name = 'Sofia Pro Soft'
font.size = Pt(120)

#### set Document/Margins
for sec in document.sections:
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = sec.page_height, sec.page_width
    sec.page_width = new_width
    sec.page_height = new_height

##### Table setting
n=i/5
nr=n*2
table = document.add_table(rows=nr, cols=5, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
d=0#type: int
while d < nr:
    set_cell_border(
        table.cell(d, 0),
        top={"val": "nil"},
        bottom={"val": "nil"},
        start={"val": "nil"},
        end={"val": "nil"},
    )
    set_cell_border(
        table.cell(d, 1),
        top={"val": "nil"},
        bottom={"val": "nil"},
        start={"val": "nil"},
        end={"val": "nil"},
    )
    set_cell_border(
        table.cell(d, 2),
        top={"val": "nil"},
        bottom={"val": "nil"},
        start={"val": "nil"},
        end={"val": "nil"},
    )
    set_cell_border(
        table.cell(d, 3),
        top={"val": "nil"},
        bottom={"val": "nil"},
        start={"val": "nil"},
        end={"val": "nil"},
    )
    set_cell_border(
        table.cell(d, 4),
        top={"val": "nil"},
        bottom={"val": "nil"},
        start={"val": "nil"},
        end={"val": "nil"},
    )

    d = d + 1

e=0
b=1
while b <= nr:
    table.cell(e, 0).height = Cm(10.88)
    table.cell(e, 0).width = Cm(4.57)
    cell1= table.cell(e,0)
    cell2= table.cell(e,2)
    cell1.merge(cell2)
    table.cell(b, 0).height = Cm(7.63)
    table.cell(b, 0).width = Cm(4.57)
    b = b + 2
    e = e + 2
table.autofit = False

#### Add table content
a = 0
b = 1
e = 0
f = 0# type: int
while b <= nr:
    table.cell(b, 0).text = items[a]
    table.cell(b, 1).text = items[a+1]
    table.cell(b, 2).text = items[a+2]
    table.cell(b, 3).text = items[a+3]
    table.cell(b, 4).text = items[a+4]

    table.cell(b, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(b, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(b, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(b, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(b, 4).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    table.cell(b, 0).paragraphs[0].paragraph_format.space_before = Pt(18)
    table.cell(b, 1).paragraphs[0].paragraph_format.space_before = Pt(18)
    table.cell(b, 2).paragraphs[0].paragraph_format.space_before = Pt(18)
    table.cell(b, 3).paragraphs[0].paragraph_format.space_before = Pt(18)
    table.cell(b, 4).paragraphs[0].paragraph_format.space_before = Pt(18)

    paragraph = table.cell(e, 0).paragraphs[0]
    paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    run = paragraph.add_run()
    run.add_picture('/Users/luzi/PycharmProjects/PHONICSCARD/phonicimages/'+items[f]+items[f+1]+items[f+2]+items[f+3]+items[f+4]+'.png',height=Inches(3.99),width=Inches(7.06))
    int(a)
    b = b + 2
    a = a + 5
    e = e + 2
    f = f + 5
document.save('five_letter_phonics.docx')

"""


four_letter_phonics.docx


"""
# read word list
file_object = open('four_letter_phonics.txt', 'r')
words = file_object.read()
items = list(words)
i = len(items)
print(items, i)
document = Document()

#### set font
style = document.styles['Normal']
font = style.font
font.name = 'Sofia Pro Soft'
font.size = Pt(120)

#### set Document/Margins
for sec in document.sections:
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = sec.page_height, sec.page_width
    sec.page_width = new_width
    sec.page_height = new_height

##### Table setting
nr=i/2
table = document.add_table(rows=nr, cols=4, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
d=0#type: int
while d < nr:
    set_cell_border(
        table.cell(d, 0),
        top={"val": "nil"},
        bottom={"val": "nil"},
        start={"val": "nil"},
        end={"val": "nil"},
    )
    set_cell_border(
        table.cell(d, 1),
        top={"val": "nil"},
        bottom={"val": "nil"},
        start={"val": "nil"},
        end={"val": "nil"},
    )
    set_cell_border(
        table.cell(d, 2),
        top={"val": "nil"},
        bottom={"val": "nil"},
        start={"val": "nil"},
        end={"val": "nil"},
    )
    set_cell_border(
        table.cell(d, 3),
        top={"val": "nil"},
        bottom={"val": "nil"},
        start={"val": "nil"},
        end={"val": "nil"},
    )

    d = d + 1

e=0
b=1  #type: int
while b <= nr:
    table.cell(e, 0).height = Cm(10.14)
    table.cell(b, 0).width = Cm(5.71)
    table.cell(b, 1).width = Cm(5.71)
    table.cell(b, 2).width = Cm(5.71)
    table.cell(b, 3).width = Cm(5.71)
    cell1= table.cell(e,0)
    cell2= table.cell(e,3)
    cell1.merge(cell2)
    table.cell(b, 0).height = Cm(4.84)
    b = b + 2
    e = e + 2

table.autofit = False
#### Add table content
a = 0
b = 1
e = 0
f = 0# type: int
while b <= nr:
    table.cell(b, 0).text = items[a]
    table.cell(b, 1).text = items[a+1]
    table.cell(b, 2).text = items[a+2]
    table.cell(b, 3).text = items[a+3]
    table.cell(b, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(b, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(b, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(b, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(b, 0).paragraphs[0].paragraph_format.space_before = Pt(18)
    table.cell(b, 1).paragraphs[0].paragraph_format.space_before = Pt(18)
    table.cell(b, 2).paragraphs[0].paragraph_format.space_before = Pt(18)
    table.cell(b, 3).paragraphs[0].paragraph_format.space_before = Pt(18)
    paragraph = table.cell(e, 0).paragraphs[0]
    paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    run = paragraph.add_run()
    run.add_picture('/Users/luzi/PycharmProjects/PHONICSCARD/phonicimages/'+items[f]+items[f+1]+items[f+2]+items[f+3]+'.png',height=Inches(3.99),width=Inches(7.06))
    int(a)
    b = b + 2
    a = a + 4
    e = e + 2
    f = f + 4
document.save('four_letter_phonics.docx')
"""


three_letter_phonics.docx


"""
# read word list
file_object = open('three_letter_phonics.txt', 'r')
words = file_object.read()
items = list(words)
i = len(items)
print(items, i)
document = Document()

#### set font
style = document.styles['Normal']
font = style.font
font.name = 'Sofia Pro Soft'
font.size = Pt(120)

#### set Document/Margins
for sec in document.sections:
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = sec.page_height, sec.page_width
    sec.page_width = new_width
    sec.page_height = new_height

##### Table setting
n=i/3
nr=n*2
table = document.add_table(rows=nr, cols=3, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
d=0#type: int
while d < nr:
    set_cell_border(
        table.cell(d, 0),
        top={"val": "nil"},
        bottom={"val": "nil"},
        start={"val": "nil"},
        end={"val": "nil"},
    )
    set_cell_border(
        table.cell(d, 1),
        top={"val": "nil"},
        bottom={"val": "nil"},
        start={"val": "nil"},
        end={"val": "nil"},
    )
    set_cell_border(
        table.cell(d, 2),
        top={"val": "nil"},
        bottom={"val": "nil"},
        start={"val": "nil"},
        end={"val": "nil"},
    )
    d = d + 1

e=0
b=1
while b <= nr:
    table.cell(e, 0).height = Cm(10.88)
    table.cell(e, 0).width = Cm(7.62)
    cell1= table.cell(e,0)
    cell2= table.cell(e,2)
    cell1.merge(cell2)
    table.cell(b, 0).height = Cm(7.63)
    table.cell(b, 0).width = Cm(7.62)
    b = b + 2
    e = e + 2
table.autofit = False

#### Add table content
a = 0
b = 1
e = 0
f = 0# type: int
while b <= nr:
    table.cell(b, 0).text = items[a]
    table.cell(b, 1).text = items[a+1]
    table.cell(b, 2).text = items[a+2]
    table.cell(b, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(b, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(b, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(b, 0).paragraphs[0].paragraph_format.space_before = Pt(18)
    table.cell(b, 1).paragraphs[0].paragraph_format.space_before = Pt(18)
    table.cell(b, 2).paragraphs[0].paragraph_format.space_before = Pt(18)
    paragraph = table.cell(e, 0).paragraphs[0]
    paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    run = paragraph.add_run()
    run.add_picture('/Users/luzi/PycharmProjects/PHONICSCARD/phonicimages/'+items[f]+items[f+1]+items[f+2]+'.png',height=Inches(3.99),width=Inches(7.06))
    int(a)
    b = b + 2
    a = a + 3
    e = e + 2
    f = f + 3
document.save('three_letter_phonics.docx')

"""


two_letter_phonics.docx


"""
# read word list
file_object = open('two_letter_phonics.txt', 'r')
words = file_object.read()
items = list(words)
i = len(items)
print(items, i)
document = Document()

#### set font
style = document.styles['Normal']
font = style.font
font.name = 'Sofia Pro Soft'
font.size = Pt(120)

#### set Document/Margins
for sec in document.sections:
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = sec.page_height, sec.page_width
    sec.page_width = new_width
    sec.page_height = new_height

##### Table setting
nr=i
table = document.add_table(rows=nr, cols=2, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
d=0#type: int
while d < nr:
    set_cell_border(
        table.cell(d, 0),
        top={"val": "nil"},
        bottom={"val": "nil"},
        start={"val": "nil"},
        end={"val": "nil"},
    )
    set_cell_border(
        table.cell(d, 1),
        top={"val": "nil"},
        bottom={"val": "nil"},
        start={"val": "nil"},
        end={"val": "nil"},
    )
    d = d + 1

e=0
b=1  #type: int
while b <= nr:
    table.cell(e, 0).height = Cm(10.14)
    table.cell(b, 0).width = Cm(11.43)
    table.cell(b, 1).width = Cm(11.43)
    cell1= table.cell(e,0)
    cell2= table.cell(e,1)
    cell1.merge(cell2)
    table.cell(b, 0).height = Cm(4.84)
    b = b + 2
    e = e + 2

table.autofit = False
#### Add table content
a = 0
b = 1
e = 0
f = 0# type: int
while b <= nr:
    table.cell(b, 0).text = items[a]
    table.cell(b, 1).text = items[a+1]
    table.cell(b, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(b, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(b, 0).paragraphs[0].paragraph_format.space_before = Pt(18)
    table.cell(b, 1).paragraphs[0].paragraph_format.space_before = Pt(18)
    paragraph = table.cell(e, 0).paragraphs[0]
    paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    run = paragraph.add_run()
    run.add_picture('/Users/luzi/PycharmProjects/PHONICSCARD/phonicimages/'+items[f]+items[f+1]+'.png',height=Inches(3.99),width=Inches(7.06))
    int(a)
    b = b + 2
    a = a + 2
    e = e + 2
    f = f + 2
document.save('two_letter_phonics.docx')

"""


2x4_phonics_slips.docx


"""
# read word list
file_object = open('2x4_phonics_slips.txt', 'r')
words = file_object.read().split(',')
items = list(words)
i = len(items)
print(items, i)
document = Document()

#### set font
style = document.styles['Normal']
font = style.font
font.name = 'Sofia Pro Soft'
font.size = Pt(120)

#### set Document/Margins
for sec in document.sections:
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)

##### Table setting
nr=i/2

table = document.add_table(rows=nr, cols=2, style='Table Grid')
for col in table.columns:
    for cell in col.cells:
        for par in cell.paragraphs:
            par.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
e = 0
while e < nr:
    table.cell(e, 0).height = Cm(5.92)
    table.cell(e, 0).width = Cm(8.26)
    table.cell(e, 1).height = Cm(5.92)
    table.cell(e, 1).width = Cm(8.26)
    e = e + 1

table.autofit = False
#### Add table content
a = 0
b = 0
while b < nr:
    table.cell(b, 0).text = items[a]
    table.cell(b, 1).text = items[a+1]
    table.cell(b, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(b, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(b, 0).paragraphs[0].paragraph_format.space_before = Pt(18)
    table.cell(b, 1).paragraphs[0].paragraph_format.space_before = Pt(18)
    b = b + 1
    a = a + 2
document.save('2x4_phonics_slips.docx')
